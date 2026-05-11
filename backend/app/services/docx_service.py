from __future__ import annotations

from datetime import datetime
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path

from app.core.config import settings


class DocxServiceUnavailableError(RuntimeError):
    pass


class DocxParseError(ValueError):
    pass


class DocxService:
    editable_sections = {
        "Introduction": "intro_message",
        "Requirements Summary": "requirements_summary",
        "Client Email Message": "personalized_message",
    }

    metadata_keys = {
        "project name": "title",
        "valid until (yyyy-mm-dd)": "valid_until",
        "tax label": "tax_label",
        "tax rate (%)": "tax_rate",
        "currency": "currency",
    }

    def _load_document_class(self):
        try:
            from docx import Document
        except ImportError as exc:
            raise DocxServiceUnavailableError(
                "DOCX generation dependencies are unavailable. Install python-docx before generating quotation drafts."
            ) from exc
        return Document

    def _resolve_logo_path(self) -> Path | None:
        asset_root = Path(__file__).resolve().parents[1] / "assets"
        candidate_paths = [
            asset_root / "muskit_logo.jpeg",
            asset_root / "muskit_logo.jpg",
            asset_root / "muskit_logo.png",
        ]
        for candidate in candidate_paths:
            if candidate.exists():
                return candidate
        return None

    def render_quotation(self, context: dict, destination: Path) -> Path:
        document_class = self._load_document_class()
        document = document_class()

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        lead = context["lead"]
        quotation = context["quotation"]
        items = context["items"]
        sections = context.get("sections") or []

        document.core_properties.title = quotation.title
        document.core_properties.subject = quotation.quotation_number
        document.core_properties.author = settings.company_name

        logo_path = self._resolve_logo_path()
        if logo_path:
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(2.8))

        heading = document.add_heading(f"{settings.company_name} quotation draft", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph(
            "Edit this Word draft, then upload the updated .docx in the admin dashboard to regenerate the reviewed PDF quotation."
        )

        contact_paragraph = document.add_paragraph()
        contact_paragraph.add_run(f"Name: {settings.admin_name}   ").bold = True
        contact_paragraph.add_run(f"Email: {settings.admin_email}   ")
        contact_paragraph.add_run(f"Phone: {settings.admin_phone}   ")
        contact_paragraph.add_run(f"Website: {settings.company_website}")

        metadata_table = document.add_table(rows=1, cols=2)
        metadata_table.style = "Table Grid"
        metadata_table.rows[0].cells[0].text = "Field"
        metadata_table.rows[0].cells[1].text = "Value"

        metadata_rows = [
            ("Project Name", quotation.title),
            ("Quotation Number", quotation.quotation_number),
            ("Issue Date", str(quotation.created_at.date() if getattr(quotation, "created_at", None) else datetime.utcnow().date())),
            ("Valid Until (YYYY-MM-DD)", str(quotation.valid_until)),
            ("Client Name", lead.full_name),
            ("Client Company", lead.company or "-"),
            ("Client Email", lead.email),
            ("Tax Label", quotation.tax_label),
            ("Tax Rate (%)", str(quotation.tax_rate)),
            ("Currency", quotation.currency),
        ]
        for label, value in metadata_rows:
            row = metadata_table.add_row().cells
            row[0].text = label
            row[1].text = value or ""

        document.add_heading("Introduction", level=1)
        document.add_paragraph(quotation.intro_message or "")

        document.add_heading("Requirements Summary", level=1)
        document.add_paragraph(quotation.requirements_summary or lead.client_requirements_text or "")

        document.add_heading("Line Items", level=1)
        items_table = document.add_table(rows=1, cols=6)
        items_table.style = "Table Grid"
        header = items_table.rows[0].cells
        header[0].text = "Sl. No."
        header[1].text = "Descriptions"
        header[2].text = "Unit"
        header[3].text = "Qty"
        header[4].text = "Unit Rate"
        header[5].text = "Amount"
        for index, item in enumerate(items, start=1):
            row = items_table.add_row().cells
            row[0].text = str(index)
            row[1].text = item.title if not item.description else f"{item.title}\n{item.description}"
            row[2].text = item.unit or "Nos"
            row[3].text = str(item.quantity)
            row[4].text = str(item.unit_price)
            row[5].text = str(item.line_total)

        document.add_heading("Additional Sections", level=1)
        sections_table = document.add_table(rows=1, cols=2)
        sections_table.style = "Table Grid"
        sections_table.rows[0].cells[0].text = "Section Title"
        sections_table.rows[0].cells[1].text = "Content"
        if sections:
            for section in sections:
                row = sections_table.add_row().cells
                row[0].text = section.get("title", "")
                row[1].text = section.get("content", "")
        else:
            row = sections_table.add_row().cells
            row[0].text = "Scope of Work"
            row[1].text = ""

        document.add_heading("Client Email Message", level=1)
        document.add_paragraph(quotation.personalized_message or "")

        document.add_paragraph(
            "Notes: you can edit line items, section titles, and section content. Keep table headers unchanged so the admin upload parser can regenerate the PDF correctly."
        )

        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)
        return destination

    def parse_quotation(self, file_bytes: bytes) -> dict:
        document_class = self._load_document_class()
        document = document_class(BytesIO(file_bytes))

        if len(document.tables) < 3:
            raise DocxParseError("The uploaded DOCX does not match the quotation draft template.")

        metadata = self._parse_metadata_table(document.tables[0])
        items = self._parse_items_table(document.tables[1])
        sections = self._parse_sections_table(document.tables[2])
        if not items:
            raise DocxParseError("At least one quotation line item is required in the uploaded DOCX.")

        fixed_sections = self._parse_fixed_sections(document)

        return {
            "title": metadata.get("title") or "Custom SaaS Proposal",
            "valid_until": self._parse_date(metadata.get("valid_until")),
            "tax_label": metadata.get("tax_label") or settings.company_tax_label,
            "tax_rate": self._parse_decimal(metadata.get("tax_rate"), field_label="Tax Rate", allow_zero=True),
            "currency": (metadata.get("currency") or settings.razorpay_currency).strip().upper(),
            "intro_message": fixed_sections.get("intro_message"),
            "requirements_summary": fixed_sections.get("requirements_summary"),
            "personalized_message": fixed_sections.get("personalized_message"),
            "items": items,
            "sections": sections,
        }

    def _parse_metadata_table(self, table) -> dict[str, str]:
        parsed: dict[str, str] = {}
        for row in table.rows[1:]:
            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()
            normalized_key = self.metadata_keys.get(key)
            if normalized_key:
                parsed[normalized_key] = value
        return parsed

    def _parse_items_table(self, table) -> list[dict]:
        parsed_items: list[dict] = []
        for index, row in enumerate(table.rows[1:], start=1):
            title_block = row.cells[1].text.strip()
            unit_text = row.cells[2].text.strip() or "Nos"
            quantity_text = row.cells[3].text.strip()
            unit_price_text = row.cells[4].text.strip()
            if not any([title_block, unit_text, quantity_text, unit_price_text]):
                continue

            title_lines = [line.strip() for line in title_block.splitlines() if line.strip()]
            title = title_lines[0] if title_lines else ""
            description = "\n".join(title_lines[1:]).strip() or None
            if not title:
                raise DocxParseError(f"Line item {index} is missing a description title in the uploaded DOCX.")

            parsed_items.append(
                {
                    "title": title,
                    "description": description,
                    "unit": unit_text,
                    "quantity": self._parse_decimal(quantity_text, field_label=f"Line item {index} quantity"),
                    "unit_price": self._parse_decimal(unit_price_text, field_label=f"Line item {index} unit price"),
                }
            )
        return parsed_items

    def _parse_sections_table(self, table) -> list[dict]:
        parsed_sections: list[dict] = []
        for row in table.rows[1:]:
            title = row.cells[0].text.strip()
            content = row.cells[1].text.strip()
            if not title and not content:
                continue
            if not title:
                raise DocxParseError("Each additional section row requires a section title in the uploaded DOCX.")
            parsed_sections.append({"title": title, "content": content})
        return parsed_sections

    def _parse_fixed_sections(self, document) -> dict[str, str | None]:
        sections = {target: [] for target in self.editable_sections.values()}
        active_section: str | None = None

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text in self.editable_sections:
                active_section = self.editable_sections[text]
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                active_section = None
                continue
            if active_section is None:
                continue
            sections[active_section].append(paragraph.text.rstrip())

        return {key: self._join_section_lines(value) for key, value in sections.items()}

    def _join_section_lines(self, lines: list[str]) -> str | None:
        cleaned = [line for line in lines if line.strip()]
        if not cleaned:
            return None
        return "\n\n".join(cleaned)

    def _parse_decimal(self, value: str | None, *, field_label: str, allow_zero: bool = False) -> Decimal:
        cleaned = (value or "").strip()
        if not cleaned:
            raise DocxParseError(f"{field_label} is required in the uploaded DOCX.")
        normalized = (
            cleaned.replace(",", "")
            .replace(settings.razorpay_currency, "")
            .replace("₹", "")
            .replace("%", "")
            .strip()
        )
        try:
            decimal_value = Decimal(normalized)
        except InvalidOperation as exc:
            raise DocxParseError(f"{field_label} must be a valid number in the uploaded DOCX.") from exc
        if allow_zero and decimal_value < 0:
            raise DocxParseError(f"{field_label} must be zero or greater in the uploaded DOCX.")
        if not allow_zero and decimal_value <= 0:
            raise DocxParseError(f"{field_label} must be greater than zero in the uploaded DOCX.")
        return decimal_value

    def _parse_date(self, value: str | None):
        raw = (value or "").strip()
        if not raw:
            return None
        try:
            return datetime.strptime(raw, "%Y-%m-%d").date()
        except ValueError as exc:
            raise DocxParseError("Valid Until must use the YYYY-MM-DD format in the uploaded DOCX.") from exc


docx_service = DocxService()
